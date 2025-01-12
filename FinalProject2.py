import streamlit as st
import google.generativeai as gemini
from docx import Document
from groq import Groq
from dotenv import load_dotenv
import tempfile
import io
import os
from streamlit_webrtc import webrtc_streamer, VideoProcessorBase
import av
import numpy as np
import wave

# Load environment variables
load_dotenv()

GROQ_API_KEY = st.secrets["groqkey"]
GEMINI_API_KEY = st.secrets["key"]

# Configure Gemini
gemini.configure(api_key=GEMINI_API_KEY)

class AudioProcessor(VideoProcessorBase):
    def __init__(self) -> None:
        self.frames = []
        self.recording = False

    def recv(self, frame: av.AudioFrame) -> av.AudioFrame:
        """Process audio frames"""
        if self.recording:
            # Convert audio frame to numpy array
            audio_data = frame.to_ndarray()
            self.frames.append(audio_data.copy())
        return frame

    def save_audio(self, filename: str) -> bool:
        """Save recorded audio to file"""
        if not self.frames or len(self.frames) == 0:
            return False

        # Convert and save to WAV
        with wave.open(filename, 'wb') as wav_file:
            wav_file.setnchannels(1)  # Mono
            wav_file.setsampwidth(2)  # 2 bytes per sample
            wav_file.setframerate(48000)  # Sample rate
            
            # Convert frames to bytes and write
            for frame in self.frames:
                wav_file.writeframes(frame.tobytes())
        
        return True

def transcribe_audio(file_path):
    """Transcribe audio using Groq"""
    client = Groq(api_key=GROQ_API_KEY)
    
    with open(file_path, "rb") as file:
        transcription = client.audio.transcriptions.create(
            file=(file_path, file.read()),
            model="whisper-large-v3-turbo",
            prompt="use full stops where necessary",
            response_format="json",
            language="en",
            temperature=0.0
        )
    return transcription.text

def generate_notes(text):
    """Generate notes using Gemini"""
    model = gemini.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content(
        "write notes with headings and numbering (and write headings on different line): " + text
    )
    return response.text

def create_formatted_doc(text):
    """Create and format a Word document"""
    doc = Document()
    paragraph = doc.add_paragraph(text)
    
    # Format text (bold and underline between **)
    for paragraph in doc.paragraphs:
        text = paragraph.text
        runs_to_add = []
        
        while "**" in text:
            start = text.find("**")
            end = text.find("**", start + 2)
            
            if end == -1:
                break
            if start > 0:
                runs_to_add.append((text[:start], False))
            runs_to_add.append((text[start + 2:end], True))
            text = text[end + 2:]
        if text:
            runs_to_add.append((text, False))
            
        for run in paragraph.runs:
            run.text = ""
        for content, is_bold in runs_to_add:
            run = paragraph.add_run(content)
            run.bold = is_bold
            run.underline = is_bold
    
    # Save to bytes buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def main():
    st.title("Audio to Notes Converter")
    
    # Initialize session state
    if 'audio_processor' not in st.session_state:
        st.session_state.audio_processor = AudioProcessor()
    if 'notes' not in st.session_state:
        st.session_state.notes = None
    if 'doc_buffer' not in st.session_state:
        st.session_state.doc_buffer = None
    
    # File name input
    file_name = st.text_input("Enter name for your file:", "my_notes")
    
    # Audio recording interface
    st.subheader("Audio Recording")

    # Create WebRTC streamer with simplified configuration
    ctx = webrtc_streamer(
        key="audio-recorder",
        mode="AUDIO_ONLY",  # Using string instead of enum
        rtc_configuration={
            "iceServers": [{"urls": ["stun:stun.l.google.com:19302"]}]
        },
        media_stream_constraints={
            "audio": True,
            "video": False
        },
        video_processor_factory=AudioProcessor,
        async_processing=True,
    )
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("Start Recording"):
            if ctx.video_processor:
                ctx.video_processor.recording = True
                ctx.video_processor.frames = []
            st.rerun()
    
    with col2:
        if st.button("Stop Recording"):
            if ctx.video_processor:
                ctx.video_processor.recording = False
                
                # Create temporary file for audio
                with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_audio:
                    if ctx.video_processor.save_audio(temp_audio.name):
                        try:
                            with st.spinner("Transcribing audio..."):
                                transcribed_text = transcribe_audio(temp_audio.name)
                                st.info("Transcription complete!")
                            
                            with st.spinner("Generating notes..."):
                                notes = generate_notes(transcribed_text)
                                st.session_state.notes = notes
                                st.info("Notes generated!")
                            
                            # Create document
                            st.session_state.doc_buffer = create_formatted_doc(notes)
                            
                        except Exception as e:
                            st.error(f"Error processing audio: {e}")
                        finally:
                            os.unlink(temp_audio.name)
                st.rerun()
    
    # Recording status indicator
    if ctx.state.playing and ctx.video_processor and ctx.video_processor.recording:
        st.markdown("🔴 **Recording in progress...**")
    
    # Display notes if available
    if st.session_state.notes:
        st.subheader("Generated Notes")
        st.text(st.session_state.notes)
        
        # Download button
        if st.session_state.doc_buffer:
            st.download_button(
                label="Download Word Document",
                data=st.session_state.doc_buffer,
                file_name=f"{file_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
