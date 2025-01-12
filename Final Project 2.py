import streamlit as st
import pyaudio
import wave
import time
import os
import google.generativeai as gemini
from docx import Document
from groq import Groq
from dotenv import load_dotenv
import tempfile
import io
import threading
import queue

# Load environment variables
load_dotenv()
GROQ_API_KEY = os.getenv("groqkey")
GEMINI_API_KEY = os.getenv("key")

# Configure Gemini
gemini.configure(api_key=GEMINI_API_KEY)

def record_audio(filename, stop_flag):
    """Record audio with manual stop control"""
    chunk = 1024
    format = pyaudio.paInt16
    channels = 2
    rate = 44100
    
    p = pyaudio.PyAudio()
    stream = p.open(
        format=format,
        channels=channels,
        rate=rate,
        input=True,
        frames_per_buffer=chunk
    )
    
    frames = []
    while not stop_flag[0]:
        try:
            data = stream.read(chunk)
            frames.append(data)
        except Exception as e:
            st.error(f"Error recording: {e}")
            break
    
    stream.stop_stream()
    stream.close()
    p.terminate()
    
    # Save to file
    wf = wave.open(filename, 'wb')
    wf.setnchannels(channels)
    wf.setsampwidth(p.get_sample_size(format))
    wf.setframerate(rate)
    wf.writeframes(b''.join(frames))
    wf.close()

def transcribe_audio(file_path):
    """Transcribe audio using Groq"""
    client = Groq(api_key=GROQ_API_KEY)
    
    with open(file_path, "rb") as file:
        transcription = client.audio.transcriptions.create(
            file=(file_path, file.read()),
            model="whisper-large-v3-turbo",
            prompt="use full stops where necessary",
            response_format="json",
            language="en",
            temperature=0.0
        )
    return transcription.text

def generate_notes(text):
    """Generate notes using Gemini"""
    model = gemini.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content(
        "write notes with headings and numbering (and write headings on different line): " + text
    )
    return response.text

def create_formatted_doc(text):
    """Create and format a Word document"""
    doc = Document()
    paragraph = doc.add_paragraph(text)
    
    # Format text (bold and underline between **)
    for paragraph in doc.paragraphs:
        text = paragraph.text
        runs_to_add = []
        
        while "**" in text:
            start = text.find("**")
            end = text.find("**", start + 2)
            
            if end == -1:
                break
            if start > 0:
                runs_to_add.append((text[:start], False))
            runs_to_add.append((text[start + 2:end], True))
            text = text[end + 2:]
        if text:
            runs_to_add.append((text, False))
            
        for run in paragraph.runs:
            run.text = ""
        for content, is_bold in runs_to_add:
            run = paragraph.add_run(content)
            run.bold = is_bold
            run.underline = is_bold
    
    # Save to bytes buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def main():
    st.title("Audio to Notes Converter")
    
    # Initialize session state
    if 'recording' not in st.session_state:
        st.session_state.recording = False
    if 'audio_file' not in st.session_state:
        st.session_state.audio_file = None
    if 'notes' not in st.session_state:
        st.session_state.notes = None
    if 'doc_buffer' not in st.session_state:
        st.session_state.doc_buffer = None
    
    # File name input
    file_name = st.text_input("Enter name for your file:", "my_notes")
    
    # Recording interface
    st.subheader("Audio Recording")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("Start Recording", disabled=st.session_state.recording):
            st.session_state.recording = True
            st.session_state.stop_flag = [False]
            
            # Create temporary file for audio
            temp_audio = tempfile.NamedTemporaryFile(delete=False, suffix='.wav')
            st.session_state.audio_file = temp_audio.name
            temp_audio.close()
            
            # Start recording in a separate thread
            recording_thread = threading.Thread(
                target=record_audio,
                args=(st.session_state.audio_file, st.session_state.stop_flag)
            )
            recording_thread.start()
            st.rerun()
    
    with col2:
        if st.button("Stop Recording", disabled=not st.session_state.recording):
            if hasattr(st.session_state, 'stop_flag'):
                st.session_state.stop_flag[0] = True
                time.sleep(0.5)  # Wait for recording to finish
                
                # Process the audio
                try:
                    with st.spinner("Transcribing audio..."):
                        transcribed_text = transcribe_audio(st.session_state.audio_file)
                        st.info("Transcription complete!")
                    
                    with st.spinner("Generating notes..."):
                        notes = generate_notes(transcribed_text)
                        st.session_state.notes = notes
                        st.info("Notes generated!")
                    
                    # Create document
                    st.session_state.doc_buffer = create_formatted_doc(notes)
                    
                except Exception as e:
                    st.error(f"Error processing audio: {e}")
                finally:
                    # Cleanup
                    if st.session_state.audio_file and os.path.exists(st.session_state.audio_file):
                        os.unlink(st.session_state.audio_file)
                    st.session_state.recording = False
                    st.rerun()
    
    # Recording status indicator
    if st.session_state.recording:
        st.markdown("🔴 **Recording in progress...**")
    
    # Display notes if available
    if st.session_state.notes:
        st.subheader("Generated Notes")
        st.text(st.session_state.notes)
        
        # Download button
        if st.session_state.doc_buffer:
            st.download_button(
                label="Download Word Document",
                data=st.session_state.doc_buffer,
                file_name=f"{file_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()